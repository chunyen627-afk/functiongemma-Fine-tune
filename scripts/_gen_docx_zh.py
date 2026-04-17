#!/usr/bin/env python3
"""產生繁體中文版 FunctionGemma_RPI5_Guide_zhTW.docx"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "FunctionGemma_RPI5_教學文件_繁中.docx"
)

doc = Document()

# ==================== 樣式設定 ====================
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'  # 微軟正黑體
# 東亞字型也要設定
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
    rpr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
rFonts.set(qn('w:ascii'), 'Microsoft JhengHei')
rFonts.set(qn('w:hAnsi'), 'Microsoft JhengHei')

style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.2


def _set_east_asia_font(run_or_style, font_name='Microsoft JhengHei'):
    """設定中文字型 (東亞字型)。"""
    element = run_or_style.element if hasattr(run_or_style, 'element') else run_or_style._element
    rpr = element.find(qn('w:rPr'))
    if rpr is None:
        rpr = element.find(qn('w:pPr'))
        if rpr is not None:
            rpr = rpr.find(qn('w:rPr'))
    if rpr is None:
        return
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


# 標題樣式
for level, (size, color_hex) in enumerate([
    (22, '1B4F72'),
    (16, '1F618D'),
    (13, '2874A6'),
    (11, '2E86C1'),
], start=1):
    hs = doc.styles[f'Heading {level}']
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color_hex)
    hs.font.name = 'Microsoft JhengHei'
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    hs.paragraph_format.space_after = Pt(6)
    _set_east_asia_font(hs)

# 程式碼樣式 (中文註解也能用，所以字型選 Consolas + 微軟正黑體)
if 'Code' not in [s.name for s in doc.styles]:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = doc.styles['Code']
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.line_spacing = 1.0

# 備註樣式
if 'Note' not in [s.name for s in doc.styles]:
    note_style = doc.styles.add_style('Note', WD_STYLE_TYPE.PARAGRAPH)
else:
    note_style = doc.styles['Note']
note_style.font.name = 'Microsoft JhengHei'
note_style.font.size = Pt(9.5)
note_style.font.italic = True
note_style.font.color.rgb = RGBColor(0x56, 0x6E, 0x7A)
note_style.paragraph_format.left_indent = Cm(0.5)
_set_east_asia_font(note_style)


# ==================== Helper ====================

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_east_asia_font(run)
    return h


def add_para(text, bold=False, style_name=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    _set_east_asia_font(run)
    return p


def add_note(text):
    p = add_para('💡 ' + text, style_name='Note')
    return p


def add_code_block(code):
    for line in code.strip('\n').split('\n'):
        p = doc.add_paragraph(style='Code')
        p.add_run(line if line else ' ')
        pPr = p._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
        pPr.append(shd)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_east_asia_font(run)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2874A6" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            _set_east_asia_font(run)
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    _set_east_asia_font(run)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    _set_east_asia_font(run)
    return p


def add_page_break():
    doc.add_page_break()


# ==================== 封面 ====================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(120)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FunctionGemma 270M\n在 Raspberry Pi 5 的部署')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor.from_string('1B4F72')
run.bold = True
_set_east_asia_font(run)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('完整教學文件 (繁體中文版)')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor.from_string('2874A6')
_set_east_asia_font(run)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run(
    '微調 / GGUF 轉換 / INT8 量化\n'
    'llama.cpp 部署 / 語音+文字雙模式對話 (羅技 C930 網路攝影機)\n'
    'system_prompt.txt 單一來源 + 真實函式執行'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
_set_east_asia_font(run)

add_page_break()

# ==================== 概述 ====================
add_heading('一、整體流程概述', 1)
add_para('本文件涵蓋完整的端到端流程：')

add_numbered('Phase A：在 Google Colab 微調 FunctionGemma 270M IT')
add_numbered('Phase B：轉換並量化為 GGUF INT8 格式 (Colab 或本機 PC)')
add_numbered('Phase C：在 RPI5 上編譯 llama.cpp')
add_numbered('Phase D：在 RPI5 上部署並運行模型')
add_numbered('Phase E：替換模型 (當你有新的微調版本時)')
add_numbered('Phase F：在 RPI5 上使用語音輸入 (羅技 C930 網路攝影機)')

add_para('架構圖：', bold=True)
add_code_block("""\
[Google Colab]                    [RPI5]
 微調 → SafeTensors                編譯 llama.cpp
     |                                |
     v                                v
 轉換為 GGUF (F16)                下載 GGUF 模型
     |                                |
     v                                v
 量化為 Q8_0                       執行推論
     |                                |
     v                                v
 上傳到 Google Drive  --------→   複製到 RPI5
                                      |
                                      v
                                  語音對話
            [C930 麥克風] → [whisper.cpp] → [文字] → [llama.cpp]""")

add_page_break()

# ==================== Phase A ====================
add_heading('Phase A：在 Google Colab 微調 FunctionGemma 270M IT', 1)

add_note('你已經有這個步驟的經驗（參考 flutter_gemma 的 Colab notebook）。此處僅列出關鍵注意事項。')

add_heading('A.1 基礎模型資訊', 2)
add_bullet('HuggingFace 模型 ID：google/functiongemma-270m-it')
add_bullet('架構：Gemma 3 270M（decoder-only transformer）')
add_bullet('Context 視窗：32K tokens')

add_heading('A.2 微調流程 (Colab)', 2)
add_para('直接使用 scripts/01_finetune_functiongemma.ipynb，這是以 flutter_gemma notebook 為基礎修改的版本，主要差異：')
add_bullet('輸出格式改為 SafeTensors（供後續 GGUF 轉換使用）')
add_bullet('移除了 TFLite 轉換步驟')
add_bullet('新增模型驗證步驟')

add_heading('A.3 訓練資料格式 (官方 FunctionGemma 規格)', 2)
add_para('FunctionGemma 使用特殊的 chat format，請「不要」使用 HuggingFace 內建的 apply_chat_template。官方格式以 <start_of_turn>developer 開頭，宣告使用 declaration:name{...} 搭配 <escape>...<escape> 包裹字串值：', bold=True)
add_code_block("""\
<start_of_turn>developer
You are a model that can do function calling with the following functions
<start_function_declaration>declaration:change_background_color{description:<escape>Changes the app background color<escape>,parameters:{properties:{color:{description:<escape>The color name (red, green, blue, yellow, purple, orange)<escape>,type:<escape>STRING<escape>}},required:[<escape>color<escape>],type:<escape>OBJECT<escape>}}<end_function_declaration>
<end_of_turn>
<start_of_turn>user
Change the background to blue
<end_of_turn>
<start_of_turn>model
<start_function_call>
{"name": "change_background_color", "arguments": {"color": "blue"}}
<end_function_call>
<end_of_turn>""")
add_note('本專案已在 rpi5/system_prompt.txt 備妥官方格式的 developer block，所有推論腳本 (voice_chat.py、function_call_client.py、llama-cli 範例) 都直接讀這個檔，不要在程式裡自行組 declaration 字串。')

add_heading('A.4 訓練完成後的輸出', 2)
add_para('訓練完成後，模型會儲存為 HuggingFace 標準格式：')
add_code_block("""\
finetuned_model/
  config.json
  model.safetensors
  tokenizer.json
  tokenizer_config.json
  special_tokens_map.json
  generation_config.json""")
add_para('重要：這個 notebook 使用完整 SFT（非 LoRA），所以輸出的就是完整權重，不需要額外 merge 步驟。', bold=True)

add_page_break()

# ==================== Phase B ====================
add_heading('Phase B：轉換並量化為 GGUF 格式', 1)

add_heading('B.1 方法一：在 Colab 中完成轉換（推薦）', 2)
add_para('使用 scripts/02_convert_to_gguf.ipynb (v3 版)，這個 notebook 會自動：')
add_numbered('Clone 並編譯 llama.cpp')
add_numbered('把 llama-quantize 與 convert_hf_to_gguf.py / gguf-py/ 一起持久化到 Google Drive (llama_cpp_tools/)')
add_numbered('將 SafeTensors 轉換為 GGUF F16 格式')
add_numbered('量化為 Q8_0 (INT8)')
add_numbered('把最終 GGUF 輸出到 Drive (functiongemma_gguf/)')
add_note('v3 的兩個關鍵修正：(1) 以 -DBUILD_SHARED_LIBS=OFF 產生「靜態」llama-quantize，避免 libllama.so.0 載入失敗；(2) 同步複製 llama.cpp 原始碼中的 gguf-py/ 到 Drive 並以 PYTHONPATH 指向它，解決新版 convert_hf_to_gguf.py 參考 gguf.MODEL_ARCH.GEMMA4 但 pip 版 gguf 沒有該屬性的錯誤。第二次執行 notebook 會偵測舊版工具並自動重編。')

add_heading('B.2 方法二：在本機 PC 完成轉換', 2)
add_code_block("""\
# 1. Clone llama.cpp
git clone https://github.com/ggml-org/llama.cpp
cd llama.cpp

# 2. 安裝 Python 依賴
pip install -r requirements.txt

# 3. 轉換為 GGUF F16（中間格式）
python convert_hf_to_gguf.py /path/to/finetuned_model \\
    --outtype f16 \\
    --outfile functiongemma-270m-finetuned-f16.gguf

# 4. 量化為 Q8_0 (INT8)
./build/bin/llama-quantize \\
    functiongemma-270m-finetuned-f16.gguf \\
    functiongemma-270m-finetuned-q8_0.gguf \\
    Q8_0""")

add_heading('B.3 一步直接轉換為 Q8_0', 2)
add_code_block("""\
python convert_hf_to_gguf.py /path/to/finetuned_model \\
    --outtype q8_0 \\
    --outfile functiongemma-270m-finetuned-q8_0.gguf""")

add_heading('B.4 量化類型選擇指南', 2)
add_table(
    ['類型', '位元', '模型大小 (約)', '品質', '建議用途'],
    [
        ['F16',    '16-bit', '~540 MB', '原始品質',  '基準 / 測試'],
        ['Q8_0',   '8-bit',  '~280 MB', '接近原始',  '推薦 - RPI5 最佳選擇'],
        ['Q6_K',   '6-bit',  '~220 MB', '很高',      '品質與大小平衡'],
        ['Q4_K_M', '4-bit',  '~170 MB', '良好',      '記憶體受限時'],
        ['Q4_0',   '4-bit',  '~150 MB', '可接受',    '最快最小'],
    ],
    col_widths=[0.7, 0.7, 1.1, 0.9, 2.0]
)
add_note('建議：270M 模型本身已經很小，Q8_0 只有 ~280MB，不需要過度壓縮。Q8_0 是品質與效能的最佳平衡點。')

add_page_break()

# ==================== Phase C ====================
add_heading('Phase C：在 RPI5 上編譯 llama.cpp', 1)

add_heading('C.1 系統準備', 2)
add_code_block("""\
# 更新系統
sudo apt update && sudo apt upgrade -y

# 安裝編譯依賴
sudo apt install -y \\
    build-essential cmake git python3-pip python3-venv \\
    libopenblas-dev pkg-config curl wget""")

add_heading('C.2 編譯 llama.cpp', 2)
add_code_block("""\
# Clone llama.cpp
cd ~
git clone https://github.com/ggml-org/llama.cpp
cd llama.cpp

# 使用 CMake 編譯（啟用 OpenBLAS 加速）
cmake -B build \\
    -DCMAKE_BUILD_TYPE=Release \\
    -DGGML_BLAS=ON \\
    -DGGML_BLAS_VENDOR=OpenBLAS

# 編譯（使用 -j3 避免 OOM，Pi 5 只有 8GB）
cmake --build build --config Release -j3""")
add_note('注意：請使用 -j3 而非 -j4，因為 -j4 在 linking 階段可能耗盡 8GB RAM。編譯大約需要 5-8 分鐘 (NVMe) 或更長時間 (microSD)。')

add_heading('C.3 驗證編譯成功', 2)
add_code_block("""\
# 確認執行檔存在
ls -la build/bin/llama-cli
ls -la build/bin/llama-server
ls -la build/bin/llama-quantize

# 測試 help
./build/bin/llama-cli --help""")

add_heading('C.4（選擇性）安裝到系統路徑', 2)
add_code_block("sudo cmake --install build")

add_page_break()

# ==================== Phase D ====================
add_heading('Phase D：在 RPI5 上部署並運行', 1)

add_heading('D.1 下載官方預訓練 GGUF 模型（用於初次測試）', 2)
add_code_block("""\
# 在 RPI5 上建立模型目錄
mkdir -p ~/llama/models

# 方法一：使用 huggingface-cli（推薦）
pip install huggingface-hub
huggingface-cli download \\
    ggml-org/functiongemma-270m-it-GGUF \\
    --local-dir ~/llama/models/functiongemma-official

# 方法二：直接下載特定量化版本
wget -O ~/llama/models/functiongemma-270m-it-q8_0.gguf \\
    "https://huggingface.co/Edge-Quant/functiongemma-270m-it-Q8_0-GGUF/\\
resolve/main/functiongemma-270m-it-q8_0.gguf\"""")

add_heading('D.2 將微調模型複製到 RPI5', 2)
add_code_block("""\
# 從 PC / Google Drive 複製 GGUF 檔到 RPI5
# 方法一：SCP
scp functiongemma-270m-finetuned-q8_0.gguf pi@<RPI5_IP>:~/llama/models/

# 方法二：從 Google Drive（若透過 Colab 轉換）
# 先在 RPI5 上安裝 rclone，或直接透過瀏覽器手動下載""")

add_heading('D.3 CLI 模式（純命令列對話）', 2)
add_para('先把 system_prompt.txt (官方格式 developer block + 所有 declaration) 複製到 RPI5 的 ~/ 下，所有測試都從這個檔組 prompt：', bold=True)
add_code_block("""\
# 確認檔案存在；開頭應為 <start_of_turn>developer
ls -la ~/llama/system_prompt.txt
head -1 ~/llama/system_prompt.txt

# 以 system_prompt.txt + user turn 組出完整 prompt
USER_MSG="把背景換成藍色"
PROMPT="$(cat ~/llama/system_prompt.txt)
<start_of_turn>user
${USER_MSG}
<end_of_turn>
<start_of_turn>model
"

~/llama/llama.cpp/build/bin/llama-cli \\
    -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 -n 256 --temp 0.1 \\
    -p "$PROMPT"

# 互動模式：請用 --system-prompt-file (不是 -f)
# 為什麼？ -f 會把整個檔案當成「第一個 user turn」丟進對話，
# 導致 developer block 被誤判成使用者輸入；而 --system-prompt-file
# 會把它當成真正的 system prompt，貫穿整個對話都有效。
# -co on 啟用彩色輸出，方便觀察輸入/輸出段。
~/llama/llama.cpp/build/bin/llama-cli \\
    -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 -co on \\
    --system-prompt-file ~/llama/system_prompt.txt""")

add_heading('D.4 Server 模式（推薦用於整合應用）', 2)
add_code_block("""\
# 啟動 HTTP API server
~/llama/llama.cpp/build/bin/llama-server \\
    -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 --host 0.0.0.0 --port 8080""")

add_para('測試 API (用 jq 安全地把多行 prompt 包成 JSON)：')
add_code_block("""\
USER_MSG="把背景換成藍色"
FULL_PROMPT="$(cat ~/llama/system_prompt.txt)
<start_of_turn>user
${USER_MSG}
<end_of_turn>
<start_of_turn>model
"
curl http://localhost:8080/completion \\
    -H "Content-Type: application/json" \\
    -d "$(jq -n --arg p "$FULL_PROMPT" \\
        '{prompt:$p, n_predict:128, temperature:0.1, stop:["<end_of_turn>"]}')"
""")

add_heading('D.5 Python 客戶端（整合到你的應用）', 2)
add_para('推薦直接使用 scripts/function_call_client.py；它會自動讀取 ~/llama/system_prompt.txt，不需要在程式內維護 function 宣告：', bold=True)
add_code_block("""\
import requests

SERVER_URL = "http://localhost:8080"

def load_system_prompt(path="~/llama/system_prompt.txt"):
    import os
    with open(os.path.expanduser(path), "r", encoding="utf-8") as f:
        return f.read().rstrip()

def call_function_gemma(user_message, system_prompt=None):
    if system_prompt is None:
        system_prompt = load_system_prompt()
    prompt = (
        f"{system_prompt}\\n"
        f"<start_of_turn>user\\n{user_message}\\n<end_of_turn>\\n"
        f"<start_of_turn>model\\n"
    )
    response = requests.post(f"{SERVER_URL}/completion", json={
        "prompt": prompt, "n_predict": 128,
        "temperature": 0.1, "stop": ["<end_of_turn>"]
    })
    return response.json()["content"]""")

add_heading('D.6 設定為 systemd 服務（開機自動啟動）', 2)
add_code_block("""\
# 使用提供的安裝腳本
sudo bash scripts/install_service.sh

# 或手動建立 systemd service：
sudo tee /etc/systemd/system/functiongemma.service << 'EOF'
[Unit]
Description=FunctionGemma LLM Server
After=network.target

[Service]
Type=simple
User=pi
WorkingDirectory=/home/pi
ExecStart=/home/pi/llama/llama.cpp/build/bin/llama-server \\
    -m /home/pi/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 --host 0.0.0.0 --port 8080
Restart=on-failure
RestartSec=5

[Install]
WantedBy=multi-user.target
EOF

sudo systemctl daemon-reload
sudo systemctl enable functiongemma
sudo systemctl start functiongemma""")

add_page_break()

# ==================== Phase E ====================
add_heading('Phase E：替換模型', 1)

add_heading('E.1 替換為新的微調模型', 2)
add_para('整個流程設計為可重複執行。當你有新的微調模型時：')
add_numbered('在 Colab 重新執行 Phase A（微調）')
add_numbered('在 Colab 重新執行 Phase B（轉換 + 量化）')
add_numbered('將新的 .gguf 檔複製到 RPI5 的 ~/llama/models/')
add_numbered('重啟服務：sudo systemctl restart functiongemma')

add_heading('E.2 切換不同模型', 2)
add_para('只需修改 -m 參數指向不同的 GGUF 檔：')
add_code_block("""\
# 官方版本
~/llama/llama.cpp/build/bin/llama-cli -m ~/llama/models/functiongemma-270m-it-q8_0.gguf ...

# 你的微調版本
~/llama/llama.cpp/build/bin/llama-cli -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf ...""")

add_page_break()

# ==================== Phase F - 語音 (C930 特化) ====================
add_heading('Phase F：RPI5 語音輸入（使用羅技 C930 網路攝影機）', 1)

add_para('本章節專門針對你使用的「羅技 C930 網路攝影機內建麥克風」做設定說明。C930 是一支 UAC（USB Audio Class）相容的 USB 裝置，RPI5 的 Linux 驅動會自動辨識，不需要額外安裝驅動程式。')
add_para('使用 whisper.cpp 在 RPI5 本地進行語音轉文字，再餵入 FunctionGemma 進行推論。完全離線運作，不需要雲端 API。', bold=True)

add_heading('F.0 羅技 C930 的特性說明', 2)
add_table(
    ['項目', '規格 / 說明'],
    [
        ['裝置類型', 'USB 網路攝影機 (UAC + UVC 相容)'],
        ['麥克風', '雙陣列立體聲麥克風 (內建)'],
        ['取樣率', '支援 16kHz / 44.1kHz / 48kHz'],
        ['聲道', '立體聲 (雙聲道)'],
        ['Linux 驅動', '無需安裝 (核心內建 snd-usb-audio)'],
        ['建議拾音距離', '30 ~ 80 公分 (C930 有內建降噪)'],
    ],
    col_widths=[1.5, 4.0]
)
add_note('注意：C930 預設是立體聲，但 Whisper 只需要單聲道。錄音時需要用 -c 1 指定單聲道，或錄完後再用 sox 轉換。')

add_heading('F.1 架構', 2)
add_code_block("""\
[羅技 C930 網路攝影機 (USB)]
         |
         v
   [arecord / sox]       錄音 (16kHz, 16-bit, 單聲道 WAV)
         |
         v
   [whisper.cpp]          語音轉文字 (本地推論, ~75-466MB 模型)
         |
         v
   [FunctionGemma]        文字轉 function call (llama.cpp)
         |
         v
   [你的應用程式]          執行對應動作 (GPIO, MQTT, HTTP...)""")

add_heading('F.2 確認 RPI5 有偵測到 C930', 2)
add_para('將 C930 接上 RPI5 的 USB 埠 (建議用藍色的 USB 3.0 埠)，然後執行：')
add_code_block("""\
# 列出所有錄音裝置
arecord -l

# 你應該會看到類似這樣的輸出：
# **** List of CAPTURE Hardware Devices ****
# card 2: C930e [Logitech Webcam C930e], device 0: USB Audio [USB Audio]
#   Subdevices: 1/1
#   Subdevice #0: subdevice #0

# 確認 USB 裝置有正確偵測
lsusb | grep -i logitech
# 輸出應該包含: 046d:0843 Logitech, Inc. Webcam C930e""")

add_note('請記住 card 號碼（上例是 card 2）。如果 C930 是 card 2，那麼裝置名稱就是 "hw:2,0"，後面的指令會用到。')

add_heading('F.3 設定 C930 為預設錄音裝置 (選擇性但強烈建議)', 2)
add_para('RPI5 可能會預設使用 HDMI 或 3.5mm 耳機孔的音訊輸入，建議把 C930 設為預設錄音裝置：')
add_code_block("""\
# 方法 1：編輯 ~/.asoundrc（永久生效）
cat > ~/.asoundrc << 'EOF'
pcm.!default {
    type asym
    capture.pcm "mic"
}
pcm.mic {
    type plug
    slave {
        pcm "hw:2,0"        # 改成你 C930 的 card 號碼
        channels 2
        rate 48000
    }
}
EOF

# 方法 2：錄音時直接指定裝置（每次都要寫）
arecord -D hw:2,0 -d 5 -r 16000 -c 1 -f S16_LE /tmp/test.wav""")

add_heading('F.4 測試 C930 錄音', 2)
add_code_block("""\
# 錄製 5 秒測試音頻（直接指定 C930）
# 請把 "2" 換成你上面看到的 C930 card 號碼
arecord -D plughw:2,0 -d 5 -r 16000 -c 1 -f S16_LE /tmp/test.wav

# 播放確認（需要接喇叭或耳機）
aplay /tmp/test.wav

# 檢查檔案大小（應該約 160KB）
ls -la /tmp/test.wav

# 如果錄出來有雜音或太小聲，調整音量：
alsamixer -c 2    # 2 是 C930 的 card 號碼
# 按 F4 切換到 Capture 模式，用方向鍵調整麥克風增益""")

add_note('注意：使用 plughw:2,0 而非 hw:2,0 可以讓 ALSA 自動做取樣率/聲道轉換，比較不會出錯。')

add_heading('F.5 安裝 whisper.cpp', 2)
add_code_block("""\
# 使用一鍵安裝腳本
chmod +x scripts/setup_whisper.sh
./scripts/setup_whisper.sh""")

add_para('腳本會自動完成：')
add_numbered('安裝音頻依賴 (ALSA, SDL2, sox)')
add_numbered('Clone 並編譯 whisper.cpp')
add_numbered('下載 Whisper 模型 (可選 tiny / base / small)')
add_numbered('測試麥克風錄音')
add_numbered('測試語音轉文字')

add_heading('F.6 Whisper 模型選擇', 2)
add_table(
    ['模型', '大小', 'RPI5 轉錄速度 (5秒音頻)', '準確度', '建議用途'],
    [
        ['tiny',  '~75 MB',  '~1-2 秒',  '基本',  '4GB RPI5 / 簡單英文指令'],
        ['base',  '~142 MB', '~2-4 秒',  '良好',  'C930 推薦 - 中英文皆可'],
        ['small', '~466 MB', '~8-15 秒', '很好',  '需要高準確度時'],
    ],
    col_widths=[0.7, 0.8, 1.6, 0.8, 2.0]
)
add_note('重要：因為 C930 是網路攝影機的麥克風，音質比專業的 USB 麥克風略差，建議至少用 base 模型。如果講中文，可以考慮用 small 模型提升準確度。')

add_heading('F.7 單獨測試 Whisper 轉錄', 2)
add_code_block("""\
# 先用 C930 錄一段 5 秒音頻
arecord -D plughw:2,0 -d 5 -r 16000 -c 1 -f S16_LE /tmp/test.wav

# 用 Whisper 轉錄（自動偵測語言）
~/llama/whisper.cpp/build/bin/whisper-cli \\
    -m ~/llama/models/whisper/ggml-base.bin \\
    -f /tmp/test.wav \\
    -l auto \\
    --no-timestamps

# 指定中文（準確度較高）
~/llama/whisper.cpp/build/bin/whisper-cli \\
    -m ~/llama/models/whisper/ggml-base.bin \\
    -f /tmp/test.wav \\
    -l zh \\
    --no-timestamps""")

add_heading('F.8 使用 Voice Chat 進行語音對話', 2)
add_para('voice_chat.py 預設讀 ~/llama/system_prompt.txt 作為官方格式 system prompt，並支援「語音 / 文字」兩種輸入模式即時切換 (擔心語音辨識失準時可立刻改用鍵盤打字)。', bold=True)
add_code_block("""\
# 把 system_prompt.txt 放到 RPI5 預設路徑 (~/llama/system_prompt.txt)
# 然後啟動 voice_chat.py

# 方式一：CLI 模式（推薦，不需啟動 server），從語音開始
python3 ~/llama/scripts/voice_chat.py --mode cli \\
    --model ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    --system-prompt ~/llama/system_prompt.txt \\
    --start-mode voice

# 方式二：Server 模式（需先啟動 llama-server）
python3 ~/llama/scripts/voice_chat.py --mode server \\
    --server-url http://localhost:8080 \\
    --system-prompt ~/llama/system_prompt.txt

# 一開始就用文字輸入 (之後還能按 v 隨時切回語音)
python3 ~/llama/scripts/voice_chat.py --start-mode text

# 使用中文語音 + small 模型 (C930 推薦)
python3 ~/llama/scripts/voice_chat.py --language zh \\
    --whisper-model ~/llama/models/whisper/ggml-small.bin

# 自動靜音偵測模式 (說完自動停)；切到 text 模式時此選項無作用
python3 ~/llama/scripts/voice_chat.py --input auto""")

add_heading('F.9 Voice Chat 操作方式 (語音 / 文字即時切換)', 2)
add_para('每一回合都會看到提示列，可透過單鍵熱鍵在語音 / 文字間切換：', bold=True)
add_table(
    ['按鍵', '行為'],
    [
        ['Enter',       '依目前模式送出：語音模式→開始錄音；文字模式→提示 You: 讓你打字'],
        ['v + Enter',   '切到語音模式 (下一回合生效)'],
        ['t + Enter',   '切到文字模式'],
        ['直接打一句話', '當場用文字送出這句 (不改變模式，最快的 fallback)'],
        ['q + Enter',   '離開 (結束時會自動重置終端機背景色)'],
    ],
    col_widths=[1.4, 4.1]
)

add_para('典型輸出 (含原始官方格式、解析結果、實際執行結果)：', bold=True)
add_code_block("""\
--- Turn 1 ---
[語音🎤] Enter=送出 / v=語音 / t=文字 / q=離開 / 或直接輸入文字 > 把背景換成藍色
  [🤔 Thinking...]
  ┌─ 原始輸出 (official format) ──────────────
  │ <start_function_call>
  │ {"name": "change_background_color", "arguments": {"color": "blue"}}
  │ <end_function_call>
  └──────────────────────────────────────────
  🔧 解析: change_background_color({"color": "blue"})
  ⚙  執行結果: [UI] ✅ 終端機背景已改為 blue""")

add_note('parse_function_call() 同時支援兩種模型輸出：純 JSON 以及 system_prompt.txt 的 declaration 風格 (name{key:<escape>value<escape>,...})。')

add_heading('F.10 內建函式真實執行效果', 2)
add_para('voice_chat.py 對 system_prompt.txt 中宣告的 4 個函式提供了「真實副作用」的 handler，不再是純字串回覆：', bold=True)
add_table(
    ['Function', '實際行為'],
    [
        ['change_background_color', 'ANSI \\033[4Xm + 清畫面，終端機背景真的換色 (支援 red/green/blue/yellow/purple/cyan/white/black/orange)'],
        ['change_app_title',        'OSC \\033]0;TITLE\\007，真的改終端機視窗標題'],
        ['show_alert',              '印出 Unicode 邊框對話方塊 + 終端機 bell'],
        ['get_current_weather',     '呼叫 wttr.in HTTP API (免 API key) 拿真實天氣'],
    ],
    col_widths=[1.8, 3.7]
)

add_para('若要新增 / 替換函式，做兩件事：', bold=True)
add_numbered('在 ~/llama/system_prompt.txt 加入新的 <start_function_declaration>...<end_function_declaration> 宣告 (官方格式)')
add_numbered('在 voice_chat.py 的 FUNCTION_HANDLERS 字典加上對應的 Python 函式')
add_code_block("""\
# voice_chat.py
def _fn_turn_on_light(args):
    room = args.get("room", "?")
    # 範例：透過 MQTT 控制智慧燈泡
    import paho.mqtt.publish as publish
    publish.single(f"home/{room}/light", "ON", hostname="localhost")
    return f"[MQTT] Light ON in {room}"

FUNCTION_HANDLERS["turn_on_light"] = _fn_turn_on_light""")
add_note('不要在腳本裡硬寫 function JSON：宣告只出現在 system_prompt.txt，腳本只負責「執行」。')

add_heading('F.11 疑難排解（C930 常見問題）', 2)

add_para('問題 1：arecord -l 找不到 C930', bold=True)
add_bullet('確認 C930 的 USB 線有插好，優先使用藍色的 USB 3.0 埠')
add_bullet('執行 dmesg | tail 查看核心訊息，確認裝置有被辨識')
add_bullet('如果用 USB Hub，改直接插入 RPI5，因為電源可能不足')

add_para('問題 2：錄音有雜音或太小聲', bold=True)
add_code_block("""\
# 調整 C930 麥克風增益 (把 "2" 換成你的 card 號碼)
alsamixer -c 2

# 或用命令列調整
amixer -c 2 sset 'Mic' 80%""")

add_para('問題 3：Whisper 轉錄不準確', bold=True)
add_bullet('使用更大的模型：ggml-base.bin → ggml-small.bin')
add_bullet('指定語言而不要用 auto：--language zh（中文）或 --language en（英文）')
add_bullet('靠近 C930 講話（30-50 公分最佳）')
add_bullet('確認錄音是 16kHz 單聲道（-r 16000 -c 1）')
add_bullet('中文建議用 small 模型（C930 非專業麥克風，小模型易誤判）')

add_para('問題 4：延遲太高', bold=True)
add_bullet('Whisper tiny 最快（~1-2 秒轉錄），但準確度較低')
add_bullet('FunctionGemma 270M 本身很快（~0.5 秒推論）')
add_bullet('主要瓶頸在 Whisper，整體延遲 ~2-5 秒 (base 模型)')
add_bullet('可考慮使用 Whisper 量化版本進一步加速')

add_para('問題 5：程式找不到預設麥克風', bold=True)
add_para('在 voice_chat.py 中，預設使用系統預設錄音裝置。如果 RPI5 預設抓到其他裝置，請：')
add_bullet('方法 1：按照 F.3 設定 ~/.asoundrc 把 C930 設為預設')
add_bullet('方法 2：在 voice_chat.py 的 AudioRecorder.record_* 方法中，將 arecord 指令加上 -D plughw:2,0 參數')

add_page_break()

# ==================== 效能預期 ====================
add_heading('二、RPI5 效能預期 (C930 + 8GB RAM)', 1)
add_table(
    ['指標', '預期值'],
    [
        ['FunctionGemma Q8_0 模型大小', '~280 MB'],
        ['Whisper base 模型大小', '~142 MB'],
        ['RAM 使用量（兩者合計）', '~500-700 MB'],
        ['FunctionGemma 推論速度', '~20-50 tokens/sec'],
        ['Whisper 轉錄速度（5秒音頻, base 模型）', '~2-4 秒'],
        ['語音指令端到端延遲', '~3-6 秒'],
        ['首次模型載入時間', '1-3 秒'],
        ['C930 拾音最佳距離', '30-50 公分'],
    ],
    col_widths=[3.5, 2.0]
)
add_note('RPI5 (Cortex-A76 @ 2.4GHz) 跑 270M 模型綽綽有餘。強烈建議使用 8GB RAM 版本的 RPI5。')

# ==================== 一般疑難排解 ====================
add_heading('三、一般疑難排解', 1)

add_heading('編譯失敗：OOM (記憶體不足)', 2)
add_code_block("# 減少並行 jobs\ncmake --build build --config Release -j2")

add_heading('模型載入失敗："unknown model architecture"', 2)
add_code_block("""\
# 確認 llama.cpp 版本夠新（需要支援 Gemma 3）
cd ~/llama/llama.cpp && git pull && cmake --build build --config Release -j3""")

add_heading('轉換失敗："missing tokenizer files"', 2)
add_para('確保 finetuned_model 資料夾包含完整的 tokenizer 檔案：')
add_code_block("""\
ls finetuned_model/tokenizer.json
ls finetuned_model/tokenizer_config.json""")

add_heading('轉換失敗："libllama.so.0: cannot open shared object file"', 2)
add_para('表示 llama_cpp_tools/ 內的 llama-quantize 是舊的 dynamic build。請執行 02_convert_to_gguf.ipynb (v3) 尾段的「強制重編」cell，或直接把 Drive 上的 llama_cpp_tools/ 整個刪掉再跑一次 notebook；v3 會自動以 -DBUILD_SHARED_LIBS=OFF 重編出靜態版本。')

add_heading('轉換失敗："MODEL_ARCH has no attribute \'GEMMA4\'"', 2)
add_para('convert_hf_to_gguf.py 版本新，但 pip 的 gguf 套件版本舊。02_convert_to_gguf.ipynb (v3) 已修正：把 llama.cpp/gguf-py/ 整個目錄一併存到 Drive，並於執行轉換時用 PYTHONPATH 指向它，同時卸載 pip 版 gguf 避免衝突。')

add_heading('推論結果不正確', 2)
add_bullet('確認 prompt 格式正確：system prompt 直接讀 ~/llama/system_prompt.txt，不要自己組 declaration')
add_bullet('降低 temperature（--temp 0.1 或 --temp 0）')
add_bullet('確認量化沒有過度壓縮（用 Q8_0 而非 Q4_0）')

add_page_break()

# ==================== 檔案結構 ====================
add_heading('四、檔案結構', 1)
add_code_block("""\
rpi5/
  README.md                                  # Markdown 版說明
  FunctionGemma_RPI5_Guide.docx              # 英文版 Word 文件
  FunctionGemma_RPI5_教學文件_繁中.docx       # 繁體中文版 Word 文件（本文件）
  RPI5_指令速查.txt                           # RPI5 端純文字指令速查 (可複製貼上)
  system_prompt.txt                           # *** 官方格式的 developer block + function 宣告；所有推論腳本都讀這個檔 ***
  scripts/
    01_finetune_functiongemma.ipynb          # Phase A: Colab 微調
    02_convert_to_gguf.ipynb                 # Phase B: Colab 轉換 + 量化 (v3: 靜態編譯 + 綁 gguf-py)
    setup_rpi5.sh                            # Phase C: RPI5 一鍵安裝 llama.cpp
    setup_whisper.sh                         # Phase F: RPI5 一鍵安裝 whisper.cpp
    install_service.sh                       # Phase D: 安裝 systemd 服務
    function_call_client.py                  # Python 客戶端範例（文字，讀取 system_prompt.txt）
    voice_chat.py                            # Phase F: 語音/文字雙模式對話客戶端 (含真實函式副作用)
    test_inference.sh                        # 推論測試腳本
  llama_cpp_tools/                            # Phase B 產出 (Colab 持久化)
    llama-quantize                            # 靜態編譯；不需 libllama.so
    convert_hf_to_gguf.py
    gguf-py/                                  # 與 convert 腳本相符的 gguf 模組 (避免 GEMMA4 錯誤)
    BUILD_INFO.txt                            # 含 build_version: v3 標記
  functiongemma_gguf/                         # Phase B 最終輸出
    functiongemma-270m-finetuned-q8_0.gguf
  sample_data/
    training_data_sample.jsonl               # 訓練資料範例""")

# 儲存
doc.save(OUTPUT_PATH)
print(f"文件已儲存：{OUTPUT_PATH}")
print(f"檔案大小：{os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
