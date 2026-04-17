#!/usr/bin/env python3
"""Generate FunctionGemma_RPI5_Guide.docx from structured content."""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           "FunctionGemma_RPI5_Guide.docx")

doc = Document()

# ==================== Styles ====================
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level, (size, color_hex) in enumerate([
    (22, '1B4F72'),  # Heading 1
    (16, '1F618D'),  # Heading 2
    (13, '2874A6'),  # Heading 3
    (11, '2E86C1'),  # Heading 4
], start=1):
    hs = doc.styles[f'Heading {level}']
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color_hex)
    hs.font.name = 'Arial'
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    hs.paragraph_format.space_after = Pt(6)

# Code style
if 'Code' not in [s.name for s in doc.styles]:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = doc.styles['Code']
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.line_spacing = 1.0

# Inline code character style
if 'InlineCode' not in [s.name for s in doc.styles]:
    ic_style = doc.styles.add_style('InlineCode', WD_STYLE_TYPE.CHARACTER)
else:
    ic_style = doc.styles['InlineCode']
ic_style.font.name = 'Consolas'
ic_style.font.size = Pt(9)
ic_style.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

# Note style
if 'Note' not in [s.name for s in doc.styles]:
    note_style = doc.styles.add_style('Note', WD_STYLE_TYPE.PARAGRAPH)
else:
    note_style = doc.styles['Note']
note_style.font.name = 'Arial'
note_style.font.size = Pt(9.5)
note_style.font.italic = True
note_style.font.color.rgb = RGBColor(0x56, 0x6E, 0x7A)
note_style.paragraph_format.left_indent = Cm(0.5)

# ==================== Helper Functions ====================

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, style_name=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    return p

def add_note(text):
    """Add a note/blockquote paragraph."""
    return add_para(text, style_name='Note')

def add_code_block(code, shading=True):
    """Add a code block with optional grey background."""
    for i, line in enumerate(code.strip().split('\n')):
        p = doc.add_paragraph(style='Code')
        p.add_run(line)
        if shading:
            pPr = p._element.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
            pPr.append(shd)
    # Small spacer after code
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)
    spacer_fmt = spacer.paragraph_format
    spacer_fmt.line_spacing = Pt(4)

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2874A6" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            # Alternate row shading
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return table

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_page_break():
    doc.add_page_break()

# ==================== Document Content ====================

# --- Title Page ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(120)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FunctionGemma 270M\non Raspberry Pi 5')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor.from_string('1B4F72')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complete Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor.from_string('2874A6')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run(
    'Fine-tuning / GGUF Conversion / INT8 Quantization\n'
    'llama.cpp Deployment / Voice Input'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_page_break()

# --- Overview ---
add_heading('Overview', 1)
add_para('This guide covers the complete pipeline:')

add_numbered('Phase A - Fine-tune FunctionGemma 270M IT (Google Colab)')
add_numbered('Phase B - Convert & Quantize to GGUF INT8 (Colab or PC)')
add_numbered('Phase C - Build llama.cpp on RPI5')
add_numbered('Phase D - Deploy & Run on RPI5')
add_numbered('Phase F - Voice Input on RPI5')

add_para('Architecture Overview:', bold=True)
add_code_block("""\
[Google Colab]                    [RPI5]
Fine-tune -> SafeTensors          Build llama.cpp
     |                                |
     v                                v
Convert to GGUF (F16)            Download GGUF
     |                                |
     v                                v
Quantize to Q8_0                 Run inference
     |                                |
     v                                v
Upload to HuggingFace/GDrive --> Copy to RPI5
                                      |
                                      v
                                 Voice Chat
                                 [Mic] -> [whisper.cpp] -> [text] -> [llama.cpp]""")

add_page_break()

# ==================== Phase A ====================
add_heading('Phase A: Fine-tune FunctionGemma 270M IT (Google Colab)', 1)

add_note('You already have experience with this step (referencing the flutter_gemma colab notebook). Key notes are provided here.')

add_heading('A.1 Base Model', 2)
add_bullet('HuggingFace ID: google/functiongemma-270m-it')
add_bullet('Architecture: Gemma 3 270M (decoder-only transformer)')
add_bullet('Context Window: 32K tokens')

add_heading('A.2 Fine-tuning Flow (Colab)', 2)
add_para('Use scripts/01_finetune_functiongemma.ipynb, which is based on the flutter_gemma notebook with the following changes:')
add_bullet('Output format changed to SafeTensors (for GGUF conversion)')
add_bullet('Removed TFLite conversion steps')
add_bullet('Added model validation step')

add_heading('A.3 Training Data Format (official FunctionGemma spec)', 2)
add_para('FunctionGemma uses a custom chat format. Do NOT use HuggingFace\'s apply_chat_template. The official format starts with <start_of_turn>developer and declares functions via declaration:name{...} using <escape>...<escape> to wrap string values:')
add_code_block("""\
<start_of_turn>developer
You are a model that can do function calling with the following functions
<start_function_declaration>declaration:change_background_color{description:<escape>Changes the app background color<escape>,parameters:{properties:{color:{description:<escape>The color name<escape>,type:<escape>STRING<escape>}},required:[<escape>color<escape>],type:<escape>OBJECT<escape>}}<end_function_declaration>
<end_of_turn>
<start_of_turn>user
Change the background to blue
<end_of_turn>
<start_of_turn>model
<start_function_call>
{"name": "change_background_color", "arguments": {"color": "blue"}}
<end_function_call>
<end_of_turn>""")
add_note('rpi5/system_prompt.txt already contains the official developer block. All inference scripts (voice_chat.py, function_call_client.py, llama-cli examples) read that file directly -- do NOT build declaration strings inside your code.')

add_heading('A.4 Training Output', 2)
add_para('After training, the model is saved in standard HuggingFace format:')
add_code_block("""\
finetuned_model/
  config.json
  model.safetensors
  tokenizer.json
  tokenizer_config.json
  special_tokens_map.json
  generation_config.json""")
add_para('Important: This notebook uses Full SFT (not LoRA), so the output contains full weights. No additional merge step is needed.', bold=True)

add_page_break()

# ==================== Phase B ====================
add_heading('Phase B: Convert & Quantize to GGUF (Colab or PC)', 1)

add_heading('B.1 Method 1: Convert in Colab (Recommended)', 2)
add_para('Use scripts/02_convert_to_gguf.ipynb (v3). This notebook will:')
add_numbered('Clone and build llama.cpp')
add_numbered('Persist llama-quantize and convert_hf_to_gguf.py / gguf-py/ to Google Drive (llama_cpp_tools/)')
add_numbered('Convert SafeTensors to GGUF F16')
add_numbered('Quantize to Q8_0 (INT8)')
add_numbered('Write final GGUF to Drive (functiongemma_gguf/)')
add_note('Two key v3 fixes: (1) builds with -DBUILD_SHARED_LIBS=OFF to produce a static llama-quantize so it never fails with libllama.so.0 missing; (2) bundles llama.cpp\'s gguf-py/ and sets PYTHONPATH to it, resolving the "MODEL_ARCH has no attribute \'GEMMA4\'" error caused by a stale pip-installed gguf. Re-running the notebook auto-detects old tools and forces a rebuild.')

add_heading('B.2 Method 2: Convert on Local PC', 2)
add_code_block("""\
# 1. Clone llama.cpp
git clone https://github.com/ggml-org/llama.cpp
cd llama.cpp

# 2. Install Python dependencies
pip install -r requirements.txt

# 3. Convert to GGUF F16 (intermediate format)
python convert_hf_to_gguf.py /path/to/finetuned_model \\
    --outtype f16 \\
    --outfile functiongemma-270m-finetuned-f16.gguf

# 4. Quantize to Q8_0 (INT8)
./build/bin/llama-quantize \\
    functiongemma-270m-finetuned-f16.gguf \\
    functiongemma-270m-finetuned-q8_0.gguf \\
    Q8_0""")

add_heading('B.3 One-step Direct Conversion', 2)
add_code_block("""\
python convert_hf_to_gguf.py /path/to/finetuned_model \\
    --outtype q8_0 \\
    --outfile functiongemma-270m-finetuned-q8_0.gguf""")

add_heading('B.4 Quantization Type Guide', 2)
add_table(
    ['Type', 'Bits', 'Size (approx)', 'Quality', 'Recommended Use'],
    [
        ['F16',    '16-bit', '~540 MB', 'Original',      'Baseline / Testing'],
        ['Q8_0',   '8-bit',  '~280 MB', 'Near-original', 'Recommended for RPI5'],
        ['Q6_K',   '6-bit',  '~220 MB', 'Very high',     'Quality/size balance'],
        ['Q4_K_M', '4-bit',  '~170 MB', 'Good',          'Memory-constrained'],
        ['Q4_0',   '4-bit',  '~150 MB', 'Acceptable',    'Fastest / smallest'],
    ],
    col_widths=[0.8, 0.7, 1.0, 1.0, 2.0]
)
add_note('Recommendation: The 270M model is already very small. Q8_0 at ~280MB provides the best quality/performance balance. No need for aggressive compression.')

add_page_break()

# ==================== Phase C ====================
add_heading('Phase C: Build llama.cpp on RPI5', 1)

add_heading('C.1 System Preparation', 2)
add_code_block("""\
# Update system
sudo apt update && sudo apt upgrade -y

# Install build dependencies
sudo apt install -y \\
    build-essential cmake git python3-pip python3-venv \\
    libopenblas-dev pkg-config curl wget""")

add_heading('C.2 Build llama.cpp', 2)
add_code_block("""\
# Clone llama.cpp
cd ~
git clone https://github.com/ggml-org/llama.cpp
cd llama.cpp

# Build with CMake (enable OpenBLAS acceleration)
cmake -B build \\
    -DCMAKE_BUILD_TYPE=Release \\
    -DGGML_BLAS=ON \\
    -DGGML_BLAS_VENDOR=OpenBLAS

# Compile (use -j3 to avoid OOM on 8GB Pi 5)
cmake --build build --config Release -j3""")
add_note('Note: Use -j3 instead of -j4. Using -j4 may exhaust 8GB RAM during the linking phase. Build takes approximately 5-8 minutes (NVMe) or longer (microSD).')

add_heading('C.3 Verify Build', 2)
add_code_block("""\
# Verify executables exist
ls -la build/bin/llama-cli
ls -la build/bin/llama-server
ls -la build/bin/llama-quantize

# Test help
./build/bin/llama-cli --help""")

add_heading('C.4 (Optional) Install to System Path', 2)
add_code_block("sudo cmake --install build")

add_page_break()

# ==================== Phase D ====================
add_heading('Phase D: Deploy & Run on RPI5', 1)

add_heading('D.1 Download Pre-trained GGUF Model (Official, for Testing)', 2)
add_code_block("""\
# Create model directory on RPI5
mkdir -p ~/llama/models

# Method 1: Using huggingface-cli (recommended)
pip install huggingface-hub
huggingface-cli download \\
    ggml-org/functiongemma-270m-it-GGUF \\
    --local-dir ~/llama/models/functiongemma-official

# Method 2: Direct download of specific quantized version
wget -O ~/llama/models/functiongemma-270m-it-q8_0.gguf \\
    "https://huggingface.co/Edge-Quant/functiongemma-270m-it-Q8_0-GGUF/\\
resolve/main/functiongemma-270m-it-q8_0.gguf\"""")

add_heading('D.2 Copy Fine-tuned Model to RPI5', 2)
add_code_block("""\
# From your PC/Google Drive, copy GGUF file to RPI5
# Method 1: SCP
scp functiongemma-270m-finetuned-q8_0.gguf pi@<RPI5_IP>:~/llama/models/

# Method 2: From Google Drive (if converted in Colab)
# Install rclone on RPI5 or download manually""")

add_heading('D.3 CLI Mode', 2)
add_code_block("""\
# Copy system_prompt.txt to ~/ first (official developer block with declarations).
# All tests build prompts from this file -- never hand-author declarations.

ls -la ~/llama/system_prompt.txt

USER_MSG="Change the background to blue"
PROMPT="$(cat ~/llama/system_prompt.txt)
<start_of_turn>user
${USER_MSG}
<end_of_turn>
<start_of_turn>model
"

~/llama/llama.cpp/build/bin/llama-cli \\
    -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 -n 256 --temp 0.1 \\
    -p "$PROMPT"

# Interactive mode: use --system-prompt-file (NOT -f)
# Why? -f treats the entire file as the "first user turn",
# so the developer block is misinterpreted as user input.
# --system-prompt-file keeps it as a persistent system prompt
# across the whole conversation.  -co on enables color output.
~/llama/llama.cpp/build/bin/llama-cli \\
    -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 -co on \\
    --system-prompt-file ~/llama/system_prompt.txt""")

add_heading('D.4 Server Mode (Recommended for App Integration)', 2)
add_code_block("""\
# Start HTTP API server
~/llama/llama.cpp/build/bin/llama-server \\
    -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 --host 0.0.0.0 --port 8080""")

add_para('Test API (use jq to safely wrap a multi-line prompt as JSON):')
add_code_block("""\
USER_MSG="Change the background to blue"
FULL_PROMPT="$(cat ~/llama/system_prompt.txt)
<start_of_turn>user
${USER_MSG}
<end_of_turn>
<start_of_turn>model
"
curl http://localhost:8080/completion \\
    -H "Content-Type: application/json" \\
    -d "$(jq -n --arg p "$FULL_PROMPT" \\
        '{prompt:$p, n_predict:128, temperature:0.1, stop:["<end_of_turn>"]}')"
""")

add_heading('D.5 Python Client (App Integration)', 2)
add_para('Prefer scripts/function_call_client.py -- it auto-loads ~/llama/system_prompt.txt so you never maintain function declarations in code:')
add_code_block("""\
import os, requests

SERVER_URL = "http://localhost:8080"

def load_system_prompt(path="~/llama/system_prompt.txt"):
    with open(os.path.expanduser(path), "r", encoding="utf-8") as f:
        return f.read().rstrip()

def call_function_gemma(user_message, system_prompt=None):
    if system_prompt is None:
        system_prompt = load_system_prompt()
    prompt = (
        f"{system_prompt}\\n"
        f"<start_of_turn>user\\n{user_message}\\n<end_of_turn>\\n"
        f"<start_of_turn>model\\n"
    )
    response = requests.post(f"{SERVER_URL}/completion", json={
        "prompt": prompt, "n_predict": 128,
        "temperature": 0.1, "stop": ["<end_of_turn>"]
    })
    return response.json()["content"]""")

add_heading('D.6 Systemd Service (Auto-start on Boot)', 2)
add_code_block("""\
# Use the provided installation script
sudo bash scripts/install_service.sh

# Or manually create systemd service:
sudo tee /etc/systemd/system/functiongemma.service << 'EOF'
[Unit]
Description=FunctionGemma LLM Server
After=network.target

[Service]
Type=simple
User=pi
WorkingDirectory=/home/pi
ExecStart=/home/pi/llama/llama.cpp/build/bin/llama-server \\
    -m /home/pi/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    -c 1024 --host 0.0.0.0 --port 8080
Restart=on-failure
RestartSec=5

[Install]
WantedBy=multi-user.target
EOF

sudo systemctl daemon-reload
sudo systemctl enable functiongemma
sudo systemctl start functiongemma""")

add_page_break()

# ==================== Phase E ====================
add_heading('Phase E: Replacing Models', 1)

add_heading('E.1 Replace with New Fine-tuned Model', 2)
add_para('The entire pipeline is designed to be repeatable. When you have a new fine-tuned model:')
add_numbered('Re-run Phase A in Colab (fine-tune)')
add_numbered('Re-run Phase B in Colab (convert + quantize)')
add_numbered('Copy the new .gguf file to ~/llama/models/ on RPI5')
add_numbered('Restart the service: sudo systemctl restart functiongemma')

add_heading('E.2 Switch Between Models', 2)
add_para('Simply change the -m parameter to point to a different GGUF file:')
add_code_block("""\
# Official version
~/llama/llama.cpp/build/bin/llama-cli -m ~/llama/models/functiongemma-270m-it-q8_0.gguf ...

# Your fine-tuned version
~/llama/llama.cpp/build/bin/llama-cli -m ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf ...""")

add_page_break()

# ==================== Phase F ====================
add_heading('Phase F: Voice Input on RPI5', 1)

add_para('Uses whisper.cpp for local speech-to-text on RPI5, feeding transcribed text into FunctionGemma. Fully offline - no cloud API required.')

add_heading('Architecture', 2)
add_code_block("""\
[USB Microphone]
      |
      v
 [arecord / sox]      Record (16kHz, 16-bit, mono WAV)
      |
      v
 [whisper.cpp]         Speech-to-text (local inference, ~75-466MB model)
      |
      v
 [FunctionGemma]       Text to function call (llama.cpp)
      |
      v
 [Your Application]    Execute actions (GPIO, MQTT, HTTP...)""")

add_heading('F.1 Hardware Requirements', 2)
add_table(
    ['Item', 'Requirement'],
    [
        ['Microphone', 'USB microphone (any UAC-compatible, noise-cancelling recommended)'],
        ['RPI5 RAM', '8GB (Whisper base + FunctionGemma Q8_0 total ~420MB)'],
    ],
    col_widths=[1.5, 4.5]
)
add_note('For 4GB RPI5, use Whisper tiny model (75MB).')

add_heading('F.2 Install whisper.cpp', 2)
add_code_block("""\
# One-click installation script
chmod +x scripts/setup_whisper.sh
./scripts/setup_whisper.sh""")
add_para('The script automatically:')
add_numbered('Installs audio dependencies (ALSA, SDL2, sox)')
add_numbered('Clones and builds whisper.cpp')
add_numbered('Downloads Whisper model (choice of tiny/base/small)')
add_numbered('Tests microphone recording')
add_numbered('Tests speech-to-text')

add_heading('F.3 Whisper Model Selection', 2)
add_table(
    ['Model', 'Size', 'RPI5 Speed (5s audio)', 'Accuracy', 'Recommended Use'],
    [
        ['tiny',  '~75 MB',  '~1-2 sec', 'Basic', '4GB RPI5 / Simple English commands'],
        ['base',  '~142 MB', '~2-4 sec', 'Good',  'Recommended - English/Chinese'],
        ['small', '~466 MB', '~8-15 sec','Very good', 'When high accuracy is needed'],
    ],
    col_widths=[0.7, 0.8, 1.4, 0.8, 2.0]
)
add_note('Whisper supports multiple languages including Chinese, Japanese, English, etc. Use --language zh to specify Chinese for improved accuracy.')

add_heading('F.4 Test Microphone', 2)
add_code_block("""\
# List recording devices
arecord -l

# Record 5-second test audio
arecord -d 5 -r 16000 -c 1 -f S16_LE /tmp/test.wav

# Playback (verify audio was recorded)
aplay /tmp/test.wav

# Transcription test
~/llama/whisper.cpp/build/bin/whisper-cli \\
    -m ~/llama/models/whisper/ggml-base.bin \\
    -f /tmp/test.wav -l auto --no-timestamps""")

add_heading('F.5 Voice Chat Interactive Mode (voice/text toggle)', 2)
add_para('voice_chat.py loads ~/llama/system_prompt.txt by default and supports on-the-fly switching between voice and text input, so you can fall back to keyboard typing whenever STT under-performs.', bold=True)
add_code_block("""\
# Mode 1: CLI mode (recommended, no server needed), start with voice
python3 ~/llama/scripts/voice_chat.py --mode cli \\
    --model ~/llama/models/functiongemma-270m-finetuned-q8_0.gguf \\
    --system-prompt ~/llama/system_prompt.txt \\
    --start-mode voice

# Mode 2: Server mode (requires llama-server running)
python3 ~/llama/scripts/voice_chat.py --mode server \\
    --server-url http://localhost:8080 \\
    --system-prompt ~/llama/system_prompt.txt

# Start in text mode (can still press 'v' to switch to voice mid-chat)
python3 ~/llama/scripts/voice_chat.py --start-mode text

# Chinese voice input + small model (recommended for C930)
python3 ~/llama/scripts/voice_chat.py --language zh \\
    --whisper-model ~/llama/models/whisper/ggml-small.bin

# Auto silence detection for voice
python3 ~/llama/scripts/voice_chat.py --input auto""")

add_heading('F.6 Voice Chat Operation (per-turn hotkeys)', 2)
add_table(
    ['Key', 'Action'],
    [
        ['Enter',       'Proceed with the current mode (voice=record / text=prompt You:)'],
        ['v + Enter',   'Switch to voice mode'],
        ['t + Enter',   'Switch to text mode'],
        ['Any text',    'Send that text right now (no mode change) -- fastest fallback'],
        ['q + Enter',   'Quit (terminal background color auto-restored on exit)'],
    ],
    col_widths=[1.4, 4.1]
)

add_para('Typical output (shows raw official format + parsed call + real effect):', bold=True)
add_code_block("""\
--- Turn 1 ---
[voice] Enter=send / v=voice / t=text / q=quit / or type text > Change the background to blue
  [Thinking...]
  +-- raw output (official format) -----------
  | <start_function_call>
  | {"name": "change_background_color", "arguments": {"color": "blue"}}
  | <end_function_call>
  +-------------------------------------------
  Parsed: change_background_color({"color": "blue"})
  Result: [UI] terminal background set to blue""")

add_note('parse_function_call() accepts both pure JSON and the declaration-style format from system_prompt.txt: name{key:<escape>value<escape>,...}')

add_heading('F.7 Built-in Real Function Handlers', 2)
add_para('voice_chat.py ships with real-effect handlers for the four functions declared in system_prompt.txt (no more string stubs):', bold=True)
add_table(
    ['Function', 'Real effect'],
    [
        ['change_background_color', 'ANSI \\033[4Xm + clear screen; terminal background really changes (red/green/blue/yellow/purple/cyan/white/black/orange)'],
        ['change_app_title',        'OSC \\033]0;TITLE\\007; terminal window title really changes'],
        ['show_alert',              'Prints a Unicode box + terminal bell'],
        ['get_current_weather',     'Calls wttr.in HTTP API (no key needed) for real weather data'],
    ],
    col_widths=[1.8, 3.7]
)

add_para('To add / replace functions, do two things:', bold=True)
add_numbered('Add a new <start_function_declaration>...<end_function_declaration> block to ~/llama/system_prompt.txt (official format)')
add_numbered('Register a Python handler in voice_chat.py\'s FUNCTION_HANDLERS dict')
add_code_block("""\
# voice_chat.py
def _fn_turn_on_light(args):
    room = args.get("room", "?")
    import paho.mqtt.publish as publish
    publish.single(f"home/{room}/light", "ON", hostname="localhost")
    return f"[MQTT] Light ON in {room}"

FUNCTION_HANDLERS["turn_on_light"] = _fn_turn_on_light""")
add_note('Never hard-code function JSON in scripts. Declarations live only in system_prompt.txt; scripts only execute.')

add_heading('F.8 Troubleshooting', 2)

add_para('No microphone sound:', bold=True)
add_code_block("""\
# Check devices
arecord -l

# Check volume (make sure Capture is not muted)
alsamixer    # Press F4 to switch to Capture, use arrows to adjust

# If using PulseAudio
pactl list sources short
pactl set-source-volume @DEFAULT_SOURCE@ 100%""")

add_para('Whisper transcription inaccurate:', bold=True)
add_bullet('Use a larger model: ggml-base.bin -> ggml-small.bin')
add_bullet('Specify language: --language zh (avoid auto)')
add_bullet('Ensure quiet environment, microphone close to mouth')
add_bullet('Ensure recording is 16kHz (-r 16000)')

add_para('Latency too high:', bold=True)
add_bullet('Whisper tiny is fastest (~1-2s transcription)')
add_bullet('FunctionGemma 270M is fast (~0.5s inference)')
add_bullet('Main bottleneck is Whisper; total latency ~2-5s (base model)')

add_page_break()

# ==================== Performance ====================
add_heading('Performance Expectations (RPI5)', 1)

add_table(
    ['Metric', 'Expected Value'],
    [
        ['FunctionGemma Q8_0 model size', '~280 MB'],
        ['Whisper base model size', '~142 MB'],
        ['RAM usage (both combined)', '~500-700 MB'],
        ['FunctionGemma inference speed', '~20-50 tokens/sec'],
        ['Whisper transcription (5s audio, base)', '~2-4 sec'],
        ['Voice command end-to-end latency', '~3-6 sec'],
        ['Initial model load time', '1-3 sec'],
    ],
    col_widths=[3.0, 2.0]
)
add_note('RPI5 (Cortex-A76 @ 2.4GHz) handles the 270M model easily. 8GB RAM version of RPI5 is recommended.')

# ==================== Troubleshooting ====================
add_heading('Troubleshooting', 1)

add_heading('Build fails: OOM', 2)
add_code_block("cmake --build build --config Release -j2")

add_heading('Model load fails: "unknown model architecture"', 2)
add_code_block("""\
# Ensure llama.cpp version supports Gemma 3
cd ~/llama/llama.cpp && git pull && cmake --build build --config Release -j3""")

add_heading('Conversion fails: "missing tokenizer files"', 2)
add_para('Ensure the finetuned_model directory contains complete tokenizer files:')
add_code_block("""\
ls finetuned_model/tokenizer.json
ls finetuned_model/tokenizer_config.json""")

add_heading('Conversion fails: "libllama.so.0: cannot open shared object file"', 2)
add_para('The llama-quantize in llama_cpp_tools/ is an old dynamic build. Run the "force rebuild" cell at the end of 02_convert_to_gguf.ipynb (v3), or delete llama_cpp_tools/ from Drive and re-run the notebook; v3 always compiles with -DBUILD_SHARED_LIBS=OFF for a static binary.')

add_heading('Conversion fails: "MODEL_ARCH has no attribute \'GEMMA4\'"', 2)
add_para('convert_hf_to_gguf.py is new but pip-installed gguf is old. 02_convert_to_gguf.ipynb (v3) fixes this by bundling llama.cpp/gguf-py/ to Drive and pointing PYTHONPATH at it, while uninstalling the pip gguf package to prevent conflict.')

add_heading('Incorrect inference results', 2)
add_bullet('Verify prompt format: always read ~/llama/system_prompt.txt, never hand-author declarations')
add_bullet('Lower temperature (--temp 0.1 or --temp 0)')
add_bullet('Verify quantization is not too aggressive (use Q8_0, not Q4_0)')

add_page_break()

# ==================== File Structure ====================
add_heading('File Structure', 1)

add_code_block("""\
rpi5/
  README.md                              # This guide (Markdown)
  FunctionGemma_RPI5_Guide.docx          # This guide (Word, English)
  FunctionGemma_RPI5_教學文件_繁中.docx    # Traditional-Chinese Word guide
  RPI5_指令速查.txt                       # Plain-text command cheat sheet for RPI5
  system_prompt.txt                      # *** Official developer block + declarations; read by all inference scripts ***
  scripts/
    01_finetune_functiongemma.ipynb       # Phase A: Colab fine-tuning
    02_convert_to_gguf.ipynb             # Phase B: Colab conversion + quantization (v3: static build + bundled gguf-py)
    setup_rpi5.sh                        # Phase C: RPI5 llama.cpp installation
    setup_whisper.sh                     # Phase F: RPI5 whisper.cpp installation
    install_service.sh                   # Phase D: systemd service installation
    function_call_client.py              # Python text client (auto-loads system_prompt.txt)
    voice_chat.py                        # Phase F: Voice/text dual-mode client with real function side-effects
    test_inference.sh                    # Inference test script
  llama_cpp_tools/                        # Phase B artifacts (persisted from Colab)
    llama-quantize                        # Static binary (no libllama.so dependency)
    convert_hf_to_gguf.py
    gguf-py/                              # Matching gguf module (fixes GEMMA4 error)
    BUILD_INFO.txt                        # Contains build_version: v3
  functiongemma_gguf/                     # Phase B final output
    functiongemma-270m-finetuned-q8_0.gguf
  sample_data/
    training_data_sample.jsonl           # Training data sample""")

# ==================== Save ====================
doc.save(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
print(f"Size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
